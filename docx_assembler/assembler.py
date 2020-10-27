#   assembler.py
#
#   Copyright (C) 2020 Aloever Dulay
#
#   This program is free software: you can redistribute it and/or modify
#   it under the terms of the GNU General Public License as published by
#   the Free Software Foundation, either version 3 of the License, or
#   (at your option) any later version.
#
#   This program is distributed in the hope that it will be useful,
#   but WITHOUT ANY WARRANTY; without even the implied warranty of
#   MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#   GNU General Public License for more details.
#
#   You should have received a copy of the GNU General Public License
#   along with this program.  If not, see <https://www.gnu.org/licenses/>.
#

import sys
import subprocess
import string

from os                             import path, rename
from docxcompose.composer           import Composer
from docx                           import Document
from pathlib                        import Path
from docx_assembler.specifications  import Specifications


try:
    from comtypes import client
    import docx2pdf
except ImportError:
    # system is linux
    client = None


root = str(path.dirname(path.realpath(__file__)))


def _docx_to_pdf_linux(doc, pdf):
    cmd = 'libreoffice --convert-to pdf'.split() + [doc]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=15)

    if path.exists(doc):
        name, ext = doc.split('.')
        rename(name + '.pdf', pdf)

    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)

def _docx_to_pdf(doc, pdf):
    doc = path.abspath(doc)
    if client is None:
        return _docx_to_pdf_linux(doc, pdf)
    docx2pdf.convert(doc, pdf)

def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def enum_documents(directory):
    files = []
    for file in Path(directory).rglob('*.docx'):
        files.append(str(file))
    files.sort()
    return files

def assemble_documents(files, output_doc, output_pdf):
    merger = Composer(Document())

    for file in files:
        merger.append(Document(file), False)
        paragraph_count = len(merger.doc.paragraphs)
        if paragraph_count > 0:
            _delete_paragraph(merger.doc.paragraphs[paragraph_count - 1])

    merger.save(output_doc)
    if output_pdf != None and output_pdf != '' and not output_pdf.isspace():
        _docx_to_pdf(output_doc, output_pdf)

def assemble(argv):
    if len(argv) > 0:
        specs = Specifications(path.abspath(argv[0]))
        assemble_documents(
            enum_documents(specs.source_dir),
            specs.doc_file_path,
            specs.pdf_file_path
        )
